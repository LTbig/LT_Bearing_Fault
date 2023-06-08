import io
import os
import subprocess
import tempfile

import pdfminer
import pdfminer.layout
import pdfminer.high_level
# import pdfminer.high_level.extract
import pdfminer.settings
from PIL import Image
from docx import Document
from docx.shared import Inches


def pdf_to_word(pdf_path, output_path):
    # 创建一个临时目录
    temp_dir = tempfile.mkdtemp()
    try:
        # 使用pdfminer解析PDF文件，并提取文本和布局信息
        with open(pdf_path, "rb") as pdf_file:
            layout = pdfminer.high_level.extract_pages(pdf_file)
            text = pdfminer.high_level.extract_text(pdf_file)

        # 创建一个空的Word文档
        document = Document()

        # 添加文本内容
        paragraphs = text.split("\n")
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

        # 处理PDF中的图片和表格
        for i, page_layout in enumerate(layout):
            for element in page_layout:
                if isinstance(element, pdfminer.layout.LTImage):
                    # 处理图片
                    image_path = os.path.join(temp_dir, f"image_{i}.png")
                    with open(image_path, "wb") as image_file:
                        image_file.write(element.stream.get_data())
                    document.add_picture(image_path, width=Inches(6))
                elif isinstance(element, pdfminer.layout.LTTextBoxHorizontal):
                    # 处理表格
                    table = document.add_table(rows=1, cols=len(element.get_text().split("\n")))
                    table.style = "Table Grid"
                    row = table.rows[0]
                    for cell, text in zip(row.cells, element.get_text().split("\n")):
                        cell.text = text

        # 保存Word文档
        document.save(output_path)

    finally:
        # 删除临时目录
        subprocess.run(["rm", "-rf", temp_dir], check=True)


# 示例代码
pdf_to_word("D:/document/论文相关/故障诊断论文/中文/基于注意力机制的滚动轴承故障诊断方法研究_李秋婷.pdf", "D:/document/论文相关/故障诊断论文/中文/基于注意力机制的滚动轴承故障诊断方法研究_李秋婷.docx")